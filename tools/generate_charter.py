#!/usr/bin/env python3
"""Generate Markdown, HTML, and DOCX project charter outputs from structured JSON input."""
from __future__ import annotations

import argparse
import html
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
except ImportError:  # pragma: no cover
    Document = None


def md_table(headers, rows):
    out = ['| ' + ' | '.join(headers) + ' |', '| ' + ' | '.join(['---']*len(headers)) + ' |']
    for row in rows:
        out.append('| ' + ' | '.join(str(x).replace('\n','<br>') for x in row) + ' |')
    return '\n'.join(out)


def bullets(items):
    return '\n'.join(f'- {item}' for item in items)


def generate_markdown(data):
    roles = [
        ['Executive Sponsor', data['sponsor'], 'Approves initiation, scope/funding tradeoffs, and material risk acceptance'],
        ['Business Owner', data['business_owner'], 'Owns business outcome, pilot adoption, and operational handoff'],
        ['Project Manager', data['project_manager'], 'Leads planning, execution coordination, issue management, and change request flow'],
        ['Approval Body', data['approval_body'], 'Authorizes initiation and approves material change requests'],
    ]
    deliv = [[d['name'], d['description']] for d in data['deliverables']]
    success = [[m['type'], m['measure'], m['baseline'], m['target'], m['owner']] for m in data['success_measures']]
    risks = [[r['risk'], r['impact'], r['mitigation']] for r in data['risks']]
    deps = [[d['dependency'], d['owner'], d['timing'], d['consequence']] for d in data['dependencies']]
    decisions = [[d['decision'], d['owner'], d['due']] for d in data['open_decisions']]
    milestones = [[m['phase'], m['target'], m['exit_criteria']] for m in data['milestones']]
    approvals = [[a['role'], a['name'], a['status']] for a in data['approvals']]
    decision_rights = '\n'.join(f'- {x}' for x in data['governance']['decision_rights'])
    return f"""# Project Charter: {data['project_name']}

| Field | Value |
|---|---|
| Project ID | {data['project_id']} |
| Version | {data['version']} |
| Date | {data['date']} |
| Executive Sponsor | {data['sponsor']} |
| Business Owner | {data['business_owner']} |
| Project Manager | {data['project_manager']} |
| Approval Body | {data['approval_body']} |

## 1. Authorization

{data['authorization']}

## 2. Purpose and Business Outcome

{data['purpose']}

**Supported business outcome:** {data['business_outcome']}

**Why now:** {data['why_now']}

## 3. Ownership and Decision Rights

{md_table(['Role','Owner','Accountability'], roles)}

## 4. Scope

### In Scope

{bullets(data['in_scope'])}

### Out of Scope

{bullets(data['out_of_scope'])}

## 5. Authorized Deliverables

{md_table(['Deliverable','Description'], deliv)}

## 6. Success Measures

{md_table(['Type','Measure','Baseline','Target','Owner'], success)}

## 7. Assumptions and Constraints

### Assumptions

{bullets(data['assumptions'])}

### Constraints

{bullets(data['constraints'])}

## 8. Risks, Dependencies, and Open Decisions

### Initial Risks

{md_table(['Risk','Impact','Mitigation Direction'], risks)}

### Critical Dependencies

{md_table(['Dependency','Owner','Timing','Consequence if Late'], deps)}

### Open Decisions

{md_table(['Decision','Owner','Due'], decisions)}

## 9. Governance Rhythm, Escalation, and Change Control

**Steering cadence:** {data['governance']['steering_cadence']}

**Core team cadence:** {data['governance']['core_team_cadence']}

**Escalation path:** {data['governance']['escalation_path']}

**Decision rights:**

{decision_rights}

**Change-control model:** {data['governance']['change_control']}

## 10. Milestone Path from Initiation to Planning and Execution

{md_table(['Phase','Target','Exit Criteria'], milestones)}

## 11. Approvals

{md_table(['Role','Name','Status'], approvals)}
"""


def generate_html(data, markdown):
    lines = markdown.splitlines()
    body = []
    in_ul = False
    in_table = False
    table = []

    def flush_ul():
        nonlocal in_ul
        if in_ul:
            body.append('</ul>')
            in_ul = False

    def flush_table():
        nonlocal in_table, table
        if in_table and table:
            body.append('<table>')
            header = table[0]
            body.append('<thead><tr>' + ''.join(f'<th>{html.escape(c.strip())}</th>' for c in header) + '</tr></thead>')
            body.append('<tbody>')
            for row in table[2:]:
                body.append('<tr>' + ''.join(f'<td>{html.escape(c.strip())}</td>' for c in row) + '</tr>')
            body.append('</tbody></table>')
        in_table = False
        table = []

    for line in lines:
        if line.startswith('|') and line.endswith('|'):
            flush_ul()
            in_table = True
            table.append([c.strip() for c in line.strip('|').split('|')])
            continue
        flush_table()
        if line.startswith('# '):
            flush_ul(); body.append(f'<h1>{html.escape(line[2:])}</h1>')
        elif line.startswith('## '):
            flush_ul(); body.append(f'<h2>{html.escape(line[3:])}</h2>')
        elif line.startswith('### '):
            flush_ul(); body.append(f'<h3>{html.escape(line[4:])}</h3>')
        elif line.startswith('- '):
            if not in_ul:
                body.append('<ul>'); in_ul = True
            body.append(f'<li>{html.escape(line[2:])}</li>')
        elif not line.strip():
            flush_ul()
        else:
            flush_ul()
            safe = html.escape(line).replace('**','')
            body.append(f'<p>{safe}</p>')
    flush_ul(); flush_table()
    css = """
    body { font-family: Arial, sans-serif; color: #1f2937; line-height: 1.45; max-width: 1080px; margin: 32px auto; padding: 0 28px; }
    h1 { color: #111827; border-bottom: 3px solid #374151; padding-bottom: 12px; }
    h2 { color: #111827; border-bottom: 1px solid #d1d5db; padding-bottom: 6px; margin-top: 30px; }
    h3 { color: #374151; margin-top: 22px; }
    table { border-collapse: collapse; width: 100%; margin: 14px 0 22px 0; font-size: 0.95rem; }
    th, td { border: 1px solid #d1d5db; padding: 8px 10px; vertical-align: top; }
    th { background: #f3f4f6; text-align: left; }
    ul { margin-top: 6px; }
    @media print { body { max-width: none; margin: 0.5in; padding: 0; } h2 { break-after: avoid; } table { page-break-inside: avoid; } }
    """
    return f"<!doctype html><html><head><meta charset='utf-8'><title>{html.escape(data['project_name'])}</title><style>{css}</style></head><body>{''.join(body)}</body></html>"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style='List Bullet')


def generate_docx(data, out_path):
    if Document is None:
        raise RuntimeError('python-docx is required. Run: pip install -r requirements.txt')
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(9)
    styles['Heading 1'].font.name = 'Arial'
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 2'].font.name = 'Arial'
    styles['Heading 2'].font.size = Pt(12)
    styles['Heading 3'].font.name = 'Arial'
    styles['Heading 3'].font.size = Pt(10.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Project Charter: {data['project_name']}")
    run.bold = True
    run.font.size = Pt(17)
    add_table(doc, ['Field','Value'], [
        ['Project ID', data['project_id']], ['Version', data['version']], ['Date', data['date']],
        ['Executive Sponsor', data['sponsor']], ['Business Owner', data['business_owner']],
        ['Project Manager', data['project_manager']], ['Approval Body', data['approval_body']]
    ])
    doc.add_heading('1. Authorization', level=1); doc.add_paragraph(data['authorization'])
    doc.add_heading('2. Purpose and Business Outcome', level=1); doc.add_paragraph(data['purpose'])
    doc.add_paragraph(f"Supported business outcome: {data['business_outcome']}")
    doc.add_paragraph(f"Why now: {data['why_now']}")
    doc.add_heading('3. Ownership and Decision Rights', level=1)
    add_table(doc, ['Role','Owner','Accountability'], [
        ['Executive Sponsor', data['sponsor'], 'Approves initiation, scope/funding tradeoffs, and material risk acceptance'],
        ['Business Owner', data['business_owner'], 'Owns business outcome, pilot adoption, and operational handoff'],
        ['Project Manager', data['project_manager'], 'Leads planning, execution coordination, issue management, and change request flow'],
        ['Approval Body', data['approval_body'], 'Authorizes initiation and approves material change requests'],
    ])
    doc.add_heading('4. Scope', level=1)
    doc.add_heading('In Scope', level=2); add_bullets(doc, data['in_scope'])
    doc.add_heading('Out of Scope', level=2); add_bullets(doc, data['out_of_scope'])
    doc.add_heading('5. Authorized Deliverables', level=1)
    add_table(doc, ['Deliverable','Description'], [[d['name'], d['description']] for d in data['deliverables']])
    doc.add_heading('6. Success Measures', level=1)
    add_table(doc, ['Type','Measure','Baseline','Target','Owner'], [[m['type'], m['measure'], m['baseline'], m['target'], m['owner']] for m in data['success_measures']])
    doc.add_heading('7. Assumptions and Constraints', level=1)
    doc.add_heading('Assumptions', level=2); add_bullets(doc, data['assumptions'])
    doc.add_heading('Constraints', level=2); add_bullets(doc, data['constraints'])
    doc.add_heading('8. Risks, Dependencies, and Open Decisions', level=1)
    doc.add_heading('Initial Risks', level=2)
    add_table(doc, ['Risk','Impact','Mitigation Direction'], [[r['risk'], r['impact'], r['mitigation']] for r in data['risks']])
    doc.add_heading('Critical Dependencies', level=2)
    add_table(doc, ['Dependency','Owner','Timing','Consequence if Late'], [[d['dependency'], d['owner'], d['timing'], d['consequence']] for d in data['dependencies']])
    doc.add_heading('Open Decisions', level=2)
    add_table(doc, ['Decision','Owner','Due'], [[d['decision'], d['owner'], d['due']] for d in data['open_decisions']])
    doc.add_heading('9. Governance Rhythm, Escalation, and Change Control', level=1)
    doc.add_paragraph(f"Steering cadence: {data['governance']['steering_cadence']}")
    doc.add_paragraph(f"Core team cadence: {data['governance']['core_team_cadence']}")
    doc.add_paragraph(f"Escalation path: {data['governance']['escalation_path']}")
    doc.add_paragraph('Decision rights:')
    add_bullets(doc, data['governance']['decision_rights'])
    doc.add_paragraph(f"Change-control model: {data['governance']['change_control']}")
    doc.add_heading('10. Milestone Path from Initiation to Planning and Execution', level=1)
    add_table(doc, ['Phase','Target','Exit Criteria'], [[m['phase'], m['target'], m['exit_criteria']] for m in data['milestones']])
    doc.add_heading('11. Approvals', level=1)
    add_table(doc, ['Role','Name','Status'], [[a['role'], a['name'], a['status']] for a in data['approvals']])
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True, help='Path to structured intake JSON')
    parser.add_argument('--out-dir', required=True, help='Scenario root containing generated_* folders')
    args = parser.parse_args()
    data = json.loads(Path(args.input).read_text(encoding='utf-8'))
    out = Path(args.out_dir)
    (out/'generated_markdown').mkdir(parents=True, exist_ok=True)
    (out/'generated_html').mkdir(parents=True, exist_ok=True)
    (out/'generated_docx').mkdir(parents=True, exist_ok=True)
    md = generate_markdown(data)
    (out/'generated_markdown/project_charter.md').write_text(md, encoding='utf-8')
    (out/'generated_html/project_charter.html').write_text(generate_html(data, md), encoding='utf-8')
    generate_docx(data, out/'generated_docx/project_charter.docx')
    print('Generated Markdown, HTML, and DOCX charter outputs.')


if __name__ == '__main__':
    main()
